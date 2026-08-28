"""
CoralTwin-DT: Automated Scientific Manuscript Formatter (DOCX & PDF)
===================================================================
Compiles manuscript.md into publication-formatted DOCX and PDF documents.

Author: CoralTwin-DT Publication Lead
License: MIT
Scientific Attribution: Resultado obtenido mediante prototipo computacional del gemelo digital.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
MANUSCRIPT_MD = os.path.join(BASE_DIR, "manuscript", "manuscript.md")
DOCX_OUT = os.path.join(BASE_DIR, "manuscript", "article.docx")
PDF_OUT = os.path.join(BASE_DIR, "manuscript", "article.pdf")


def generate_docx():
    print(f"Generating formatted DOCX: {DOCX_OUT}...")
    doc = Document()

    # Configure Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    with open(MANUSCRIPT_MD, "r", encoding="utf-8") as f:
        lines = f.readlines()

    for line in lines:
        line_str = line.strip()
        if not line_str:
            continue

        if line_str.startswith("# "):
            p = doc.add_paragraph()
            run = p.add_run(line_str[2:])
            run.font.name = "Calibri"
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(27, 73, 101)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(8)
        elif line_str.startswith("## "):
            p = doc.add_paragraph()
            run = p.add_run(line_str[3:])
            run.font.name = "Calibri"
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(43, 108, 176)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
        elif line_str.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(line_str[4:])
            run.font.name = "Calibri"
            run.font.size = Pt(12)
            run.font.bold = True
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
        elif line_str.startswith("- ") or line_str.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(line_str[2:])
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
        elif line_str.startswith("1. ") or line_str.startswith("2. ") or line_str.startswith("3. ") or line_str.startswith("4. "):
            p = doc.add_paragraph()
            run = p.add_run(line_str)
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
        elif line_str.startswith("|"):
            # Render Markdown table line as code/monospace or plain text
            p = doc.add_paragraph()
            run = p.add_run(line_str)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        else:
            p = doc.add_paragraph()
            run = p.add_run(line_str)
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(4)

    doc.save(DOCX_OUT)
    print("DOCX successfully generated.")


class PDFReport(FPDF):
    def header(self):
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(100, 100, 100)
        self.cell(0, 8, "Global Change Biology | CoralTwin-DT Research Consortium", 0, 1, "R")

    def footer(self):
        self.set_y(-15)
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(100, 100, 100)
        self.cell(0, 10, f"Page {self.page_no()}", 0, 0, "C")


def generate_pdf():
    print(f"Generating formatted PDF: {PDF_OUT}...")
    pdf = PDFReport()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", "", 10)

    with open(MANUSCRIPT_MD, "r", encoding="utf-8") as f:
        lines = f.readlines()

    for line in lines:
        line_str = line.strip().encode("latin-1", "replace").decode("latin-1")
        if not line_str:
            pdf.ln(2)
            continue

        if line_str.startswith("# "):
            pdf.set_font("Helvetica", "B", 14)
            pdf.set_text_color(27, 73, 101)
            pdf.multi_cell(0, 7, line_str[2:])
            pdf.ln(2)
        elif line_str.startswith("## "):
            pdf.set_font("Helvetica", "B", 11)
            pdf.set_text_color(43, 108, 176)
            pdf.multi_cell(0, 6, line_str[3:])
            pdf.ln(1)
        elif line_str.startswith("### "):
            pdf.set_font("Helvetica", "B", 10)
            pdf.set_text_color(0, 0, 0)
            pdf.multi_cell(0, 5, line_str[4:])
            pdf.ln(1)
        else:
            pdf.set_font("Helvetica", "", 9)
            pdf.set_text_color(40, 40, 40)
            pdf.multi_cell(0, 4.5, line_str)
            pdf.ln(1)

    pdf.output(PDF_OUT)
    print("PDF successfully generated.")


if __name__ == "__main__":
    generate_docx()
    generate_pdf()
