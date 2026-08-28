"""
CoralTwin-DT: Final Submission Package Compiler (.docx / .pdf / .png)
====================================================================
Compiles all final publication assets for Ecological Informatics.

Author: CoralTwin-DT Publication Team
License: MIT
"""

import os
import shutil
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
from fpdf.enums import XPos, YPos

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.dirname(os.path.dirname(BASE_DIR))

MANUSCRIPT_MD = os.path.join(BASE_DIR, "manuscript_final.md")
SUPP_MD = os.path.join(BASE_DIR, "supplementary_material.md")
DOCX_OUT = os.path.join(BASE_DIR, "manuscript.docx")
PDF_OUT = os.path.join(BASE_DIR, "manuscript.pdf")
SUPP_PDF_OUT = os.path.join(BASE_DIR, "supplementary_material.pdf")
GA_PNG_SRC = os.path.join(PROJECT_ROOT, "10_Publication", "final_manuscript", "graphical_abstract.png")
GA_PNG_DEST = os.path.join(BASE_DIR, "graphical_abstract.png")


def copy_graphical_abstract():
    if os.path.exists(GA_PNG_SRC):
        shutil.copyfile(GA_PNG_SRC, GA_PNG_DEST)
        print(f"Copied Graphical Abstract to: {GA_PNG_DEST}")


def build_manuscript_docx():
    print(f"Compiling Manuscript DOCX: {DOCX_OUT}...")
    doc = Document()

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("CoralTwin-DT: A Cyber-Physical Digital Twin Coupling Machine Learning and Dynamical Biophysics for Coral Reef Restoration Prioritization under Thermal Stress and Ocean Acidification")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(16)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(26, 54, 93)

    p_auth = doc.add_paragraph()
    p_auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_auth = p_auth.add_run("CoralTwin-DT Doctoral Research Consortium\nTarget Journal: Ecological Informatics (Elsevier Q1)\nRepository: https://github.com/HrSly11/CoralTwin-DT.git")
    run_auth.font.name = "Arial"
    run_auth.font.size = Pt(10)
    run_auth.font.italic = True

    with open(MANUSCRIPT_MD, "r", encoding="utf-8") as f:
        lines = f.readlines()

    for line in lines:
        line_s = line.strip()
        if not line_s or line_s.startswith("```") or line_s == "---" or line_s.startswith("|"):
            continue

        if line_s.startswith("## "):
            h = doc.add_heading(line_s[3:], level=1)
            h.style.font.name = "Arial"
            h.style.font.color.rgb = RGBColor(43, 108, 176)
        elif line_s.startswith("### "):
            h = doc.add_heading(line_s[4:], level=2)
            h.style.font.name = "Arial"
            h.style.font.color.rgb = RGBColor(45, 55, 72)
        elif line_s.startswith("- "):
            p = doc.add_paragraph(line_s[2:], style="List Bullet")
            p.style.font.name = "Calibri"
            p.style.font.size = Pt(10.5)
        else:
            p = doc.add_paragraph(line_s)
            p.style.font.name = "Calibri"
            p.style.font.size = Pt(10.5)

    doc.save(DOCX_OUT)
    print("Manuscript DOCX generated.")


class EcologicalInformaticsPDF(FPDF):
    def header(self):
        self.set_font("Helvetica", "B", 8)
        self.set_text_color(100, 116, 139)
        self.cell(180, 6, "Ecological Informatics (Scopus Q1) - Final Submission Package", border=0, align="L",
                  new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        self.set_draw_color(226, 232, 240)
        self.line(15, 14, 195, 14)
        self.ln(2)

    def footer(self):
        self.set_y(-12)
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(148, 163, 184)
        self.cell(180, 8, f"Page {self.page_no()}", border=0, align="C",
                  new_x=XPos.RIGHT, new_y=YPos.TOP)


def build_manuscript_pdf():
    print(f"Compiling Manuscript PDF: {PDF_OUT}...")
    pdf = EcologicalInformaticsPDF()
    pdf.set_margins(15, 15, 15)
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    pdf.set_font("Helvetica", "B", 13)
    pdf.set_text_color(26, 54, 93)
    pdf.multi_cell(180, 6, "CoralTwin-DT: A Cyber-Physical Digital Twin Coupling Machine Learning and Dynamical Biophysics for Coral Reef Restoration Prioritization under Thermal Stress and Ocean Acidification", align="C")
    pdf.ln(3)

    pdf.set_font("Helvetica", "I", 9)
    pdf.set_text_color(71, 85, 105)
    pdf.cell(180, 5, "CoralTwin-DT Doctoral Research Consortium - Ecological Informatics (Scopus Q1)", align="C",
             new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(4)

    with open(MANUSCRIPT_MD, "r", encoding="utf-8") as f:
        lines = f.readlines()

    for line in lines:
        line_s = line.strip()
        if not line_s or line_s.startswith("```") or line_s == "---" or line_s.startswith("|"):
            continue

        if line_s.startswith("## "):
            pdf.ln(3)
            pdf.set_font("Helvetica", "B", 10.5)
            pdf.set_text_color(43, 108, 176)
            clean_t = line_s[3:].encode("latin-1", "replace").decode("latin-1")
            pdf.cell(180, 6, clean_t, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.ln(1)
        elif line_s.startswith("### "):
            pdf.ln(2)
            pdf.set_font("Helvetica", "B", 9.5)
            pdf.set_text_color(30, 41, 59)
            clean_t = line_s[4:].encode("latin-1", "replace").decode("latin-1")
            pdf.cell(180, 5, clean_t, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        elif line_s.startswith("- "):
            pdf.set_font("Helvetica", "", 8.5)
            pdf.set_text_color(51, 65, 85)
            clean_t = ("  * " + line_s[2:]).encode("latin-1", "replace").decode("latin-1")
            pdf.multi_cell(180, 4.5, clean_t)
        else:
            pdf.set_font("Helvetica", "", 8.5)
            pdf.set_text_color(51, 65, 85)
            clean_t = line_s.encode("latin-1", "replace").decode("latin-1")
            pdf.multi_cell(180, 4.5, clean_t)
            pdf.ln(1)

    pdf.output(PDF_OUT)
    print("Manuscript PDF generated.")


def build_supplementary_pdf():
    print(f"Compiling Supplementary Material PDF: {SUPP_PDF_OUT}...")
    pdf = EcologicalInformaticsPDF()
    pdf.set_margins(15, 15, 15)
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    pdf.set_font("Helvetica", "B", 13)
    pdf.set_text_color(26, 54, 93)
    pdf.multi_cell(180, 6, "Supplementary Material: CoralTwin-DT Methods and Tables", align="C")
    pdf.ln(3)

    pdf.set_font("Helvetica", "I", 9)
    pdf.set_text_color(71, 85, 105)
    pdf.cell(180, 5, "Target Journal: Ecological Informatics - Doctoral Research Consortium", align="C",
             new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(4)

    with open(SUPP_MD, "r", encoding="utf-8") as f:
        lines = f.readlines()

    for line in lines:
        line_s = line.strip()
        if not line_s or line_s.startswith("```") or line_s == "---" or line_s.startswith("|"):
            continue

        if line_s.startswith("## "):
            pdf.ln(3)
            pdf.set_font("Helvetica", "B", 10.5)
            pdf.set_text_color(43, 108, 176)
            clean_t = line_s[3:].encode("latin-1", "replace").decode("latin-1")
            pdf.cell(180, 6, clean_t, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.ln(1)
        elif line_s.startswith("### "):
            pdf.ln(2)
            pdf.set_font("Helvetica", "B", 9.5)
            pdf.set_text_color(30, 41, 59)
            clean_t = line_s[4:].encode("latin-1", "replace").decode("latin-1")
            pdf.cell(180, 5, clean_t, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        elif line_s.startswith("- "):
            pdf.set_font("Helvetica", "", 8.5)
            pdf.set_text_color(51, 65, 85)
            clean_t = ("  * " + line_s[2:]).encode("latin-1", "replace").decode("latin-1")
            pdf.multi_cell(180, 4.5, clean_t)
        else:
            pdf.set_font("Helvetica", "", 8.5)
            pdf.set_text_color(51, 65, 85)
            clean_t = line_s.encode("latin-1", "replace").decode("latin-1")
            pdf.multi_cell(180, 4.5, clean_t)
            pdf.ln(1)

    pdf.output(SUPP_PDF_OUT)
    print("Supplementary Material PDF generated.")


if __name__ == "__main__":
    copy_graphical_abstract()
    build_manuscript_docx()
    build_manuscript_pdf()
    build_supplementary_pdf()
    print("Final submission package compiled successfully.")
