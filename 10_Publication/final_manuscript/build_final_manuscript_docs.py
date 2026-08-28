"""
CoralTwin-DT: Final Publication Documents Generator (.docx and .pdf)
===================================================================
Generates the publication-ready DOCX and PDF documents for Scopus Q1 submission.

Author: CoralTwin-DT Publication Team
License: MIT
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
from fpdf.enums import XPos, YPos

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
MD_PATH = os.path.join(BASE_DIR, "manuscript_q1_final.md")
DOCX_PATH = os.path.join(BASE_DIR, "manuscript_q1.docx")
PDF_PATH = os.path.join(BASE_DIR, "manuscript_q1.pdf")


def build_docx():
    print(f"Building DOCX: {DOCX_PATH}...")
    doc = Document()

    # Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("CoralTwin-DT: A Cyber-Physical Digital Twin Coupling Machine Learning and Dynamical Biophysics for Coral Reef Restoration Prioritization under Thermal Stress and Ocean Acidification")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(16)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(26, 54, 93)

    # Authors
    p_auth = doc.add_paragraph()
    p_auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_auth = p_auth.add_run("CoralTwin-DT Doctoral Research Consortium\nComputational Oceanography & Ecological Informatics Laboratory\nTarget: Ecological Informatics (Elsevier Q1)")
    run_auth.font.name = "Arial"
    run_auth.font.size = Pt(10.5)
    run_auth.font.italic = True

    # Parse markdown lines
    with open(MD_PATH, "r", encoding="utf-8") as f:
        lines = f.readlines()

    for line in lines:
        line_s = line.strip()
        if not line_s or line_s.startswith("```") or line_s == "---":
            continue

        if line_s.startswith("## "):
            h = doc.add_heading(line_s[3:], level=1)
            h.style.font.name = "Arial"
            h.style.font.color.rgb = RGBColor(43, 108, 176)
        elif line_s.startswith("### "):
            h = doc.add_heading(line_s[4:], level=2)
            h.style.font.name = "Arial"
            h.style.font.color.rgb = RGBColor(45, 55, 72)
        elif line_s.startswith("- "):
            p = doc.add_paragraph(line_s[2:], style="List Bullet")
            p.style.font.name = "Calibri"
            p.style.font.size = Pt(10.5)
        elif line_s.startswith("|"):
            continue
        else:
            p = doc.add_paragraph(line_s)
            p.style.font.name = "Calibri"
            p.style.font.size = Pt(10.5)

    doc.save(DOCX_PATH)
    print("DOCX built successfully.")


class Q1ManuscriptPDF(FPDF):
    def header(self):
        self.set_font("Helvetica", "B", 8)
        self.set_text_color(100, 116, 139)
        self.cell(0, 7, "Ecological Informatics (Scopus Q1) - CoralTwin-DT Final Manuscript", border=0, align="L",
                  new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        self.set_draw_color(226, 232, 240)
        self.line(12, 14, 198, 14)
        self.ln(3)

    def footer(self):
        self.set_y(-12)
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(148, 163, 184)
        self.cell(0, 8, f"Page {self.page_no()}", border=0, align="C",
                  new_x=XPos.RIGHT, new_y=YPos.TOP)


def build_pdf():
    print(f"Building PDF: {PDF_PATH}...")
    pdf = Q1ManuscriptPDF()
    pdf.set_margins(15, 15, 15)
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    # Title
    pdf.set_font("Helvetica", "B", 13)
    pdf.set_text_color(26, 54, 93)
    pdf.multi_cell(180, 6, "CoralTwin-DT: A Cyber-Physical Digital Twin Coupling Machine Learning and Dynamical Biophysics for Coral Reef Restoration Prioritization under Thermal Stress and Ocean Acidification", align="C")
    pdf.ln(3)

    pdf.set_font("Helvetica", "I", 9)
    pdf.set_text_color(71, 85, 105)
    pdf.cell(180, 5, "CoralTwin-DT Doctoral Research Consortium - Target: Ecological Informatics (Scopus Q1)", align="C",
             new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(4)

    with open(MD_PATH, "r", encoding="utf-8") as f:
        lines = f.readlines()

    for line in lines:
        line_s = line.strip()
        if not line_s or line_s.startswith("```") or line_s == "---" or line_s.startswith("|"):
            continue

        if line_s.startswith("## "):
            pdf.ln(3)
            pdf.set_font("Helvetica", "B", 10.5)
            pdf.set_text_color(43, 108, 176)
            clean_t = line_s[3:].encode("latin-1", "replace").decode("latin-1")
            pdf.cell(180, 6, clean_t, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.ln(1)
        elif line_s.startswith("### "):
            pdf.ln(2)
            pdf.set_font("Helvetica", "B", 9.5)
            pdf.set_text_color(30, 41, 59)
            clean_t = line_s[4:].encode("latin-1", "replace").decode("latin-1")
            pdf.cell(180, 5, clean_t, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        elif line_s.startswith("- "):
            pdf.set_font("Helvetica", "", 8.5)
            pdf.set_text_color(51, 65, 85)
            clean_t = ("  * " + line_s[2:]).encode("latin-1", "replace").decode("latin-1")
            pdf.multi_cell(180, 4.5, clean_t)
        else:
            pdf.set_font("Helvetica", "", 8.5)
            pdf.set_text_color(51, 65, 85)
            clean_t = line_s.encode("latin-1", "replace").decode("latin-1")
            pdf.multi_cell(180, 4.5, clean_t)
            pdf.ln(1)

    pdf.output(PDF_PATH)
    print("PDF built successfully.")


if __name__ == "__main__":
    build_docx()
    build_pdf()
